import streamlit as st
import streamlit.components.v1 as components
import markdown
import re
import requests
from io import BytesIO
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import zipfile

# Page configuration
st.set_page_config(
    page_title="Markdown to PDF Converter",
    page_icon="📄",
    layout="wide"
)

# Page sizes in mm
PAGE_SIZES = {
    "A4": (210, 297),
    "A5": (148, 210),
    "Letter": (216, 279),
    "Legal": (216, 356),
}

# Theme configurations
THEMES = {
    "Professionnel": {
        "font_family": "Helvetica, Arial, sans-serif",
        "heading_color": "#1a1a2e",
        "text_color": "#333333",
        "link_color": "#0066cc",
        "code_bg": "#f4f4f4",
        "blockquote_border": "#3498db",
        "table_header_bg": "#2c3e50",
        "table_header_color": "#ffffff",
    },
    "Académique": {
        "font_family": "Georgia, Times, serif",
        "heading_color": "#2c3e50",
        "text_color": "#2c3e50",
        "link_color": "#8e44ad",
        "code_bg": "#ecf0f1",
        "blockquote_border": "#95a5a6",
        "table_header_bg": "#34495e",
        "table_header_color": "#ffffff",
    },
    "Minimaliste": {
        "font_family": "Helvetica, Arial, sans-serif",
        "heading_color": "#000000",
        "text_color": "#444444",
        "link_color": "#000000",
        "code_bg": "#fafafa",
        "blockquote_border": "#cccccc",
        "table_header_bg": "#f5f5f5",
        "table_header_color": "#333333",
    },
    "Moderne": {
        "font_family": "Segoe UI, Roboto, sans-serif",
        "heading_color": "#6c5ce7",
        "text_color": "#2d3436",
        "link_color": "#0984e3",
        "code_bg": "#dfe6e9",
        "blockquote_border": "#6c5ce7",
        "table_header_bg": "#6c5ce7",
        "table_header_color": "#ffffff",
    },
}

# Markdown extensions list - removed codehilite to fix [object Object] issue
MD_EXTENSIONS = ['extra', 'tables', 'fenced_code', 'nl2br', 'toc']

# Markdown cheatsheet
MARKDOWN_CHEATSHEET = """
### 📝 Syntaxe Markdown

| Élément | Syntaxe |
|---------|---------|
| **Gras** | `**texte**` |
| *Italique* | `*texte*` |
| ~~Barré~~ | `~~texte~~` |
| `Code` | `` `code` `` |
| [Lien](url) | `[texte](url)` |
| Image | `![alt](url)` |

### Titres
```
# H1
## H2
### H3
```

### Listes
```
- Item 1
- Item 2

1. Premier
2. Deuxième
```

### Code block
````
```python
code ici
```
````

### Citation
```
> Citation
```

### Tableau
```
| Col1 | Col2 |
|------|------|
| A    | B    |
```

### Ligne horizontale
```
---
```
"""


def generate_css(theme: dict, page_size: str, orientation: str, margins: dict, 
                 show_page_numbers: bool, header_text: str, footer_text: str,
                 watermark: str = "") -> str:
    """Generate custom CSS based on theme and page settings."""
    width, height = PAGE_SIZES[page_size]
    if orientation == "Paysage":
        width, height = height, width
    
    watermark_css = ""
    if watermark:
        watermark_css = f"""
        body::before {{
            content: "{watermark}";
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 72pt;
            color: rgba(200, 200, 200, 0.3);
            z-index: -1;
            white-space: nowrap;
        }}
        """
    
    header_content = f'content: "{header_text}";' if header_text else 'content: "";'
    footer_content = f'content: "{footer_text}";' if footer_text else ""
    page_number_css = 'content: "Page " counter(page);' if show_page_numbers else ""
    
    return f"""
@page {{
    size: {width}mm {height}mm;
    margin: {margins['top']}cm {margins['right']}cm {margins['bottom']}cm {margins['left']}cm;
    @top-center {{
        {header_content}
        font-size: 9pt;
        color: #888;
    }}
    @bottom-right {{
        {page_number_css}
        font-size: 9pt;
        color: #888;
    }}
    @bottom-left {{
        {footer_content}
        font-size: 9pt;
        color: #888;
    }}
}}
{watermark_css}
body {{
    font-family: {theme['font_family']};
    font-size: 11pt;
    line-height: 1.6;
    color: {theme['text_color']};
}}
h1, h2, h3, h4, h5, h6 {{
    margin-top: 20pt;
    margin-bottom: 12pt;
    font-weight: bold;
    line-height: 1.25;
    color: {theme['heading_color']};
}}
h1 {{ font-size: 22pt; border-bottom: 2px solid {theme['heading_color']}; padding-bottom: 6pt; }}
h2 {{ font-size: 18pt; border-bottom: 1px solid {theme['heading_color']}; padding-bottom: 4pt; }}
h3 {{ font-size: 14pt; }}
h4 {{ font-size: 12pt; }}
code {{
    background-color: {theme['code_bg']};
    padding: 2pt 4pt;
    font-family: Courier, monospace;
    font-size: 10pt;
    border-radius: 3pt;
}}
pre {{
    background-color: {theme['code_bg']};
    padding: 12pt;
    border: 1px solid #ddd;
    border-radius: 5pt;
    overflow: hidden;
    white-space: pre-wrap;
    word-wrap: break-word;
}}
pre code {{
    background-color: transparent;
    padding: 0;
}}
blockquote {{
    padding: 8pt 16pt;
    color: #555;
    border-left: 4pt solid {theme['blockquote_border']};
    margin: 12pt 0;
    background-color: #fafafa;
    font-style: italic;
}}
table {{
    border-collapse: collapse;
    width: 100%;
    margin: 12pt 0;
}}
th, td {{
    padding: 8pt 10pt;
    border: 1px solid #ddd;
    text-align: left;
}}
th {{
    background-color: {theme['table_header_bg']};
    color: {theme['table_header_color']};
    font-weight: bold;
}}
tr:nth-child(even) {{
    background-color: #fafafa;
}}
a {{
    color: {theme['link_color']};
    text-decoration: none;
}}
a:hover {{
    text-decoration: underline;
}}
ul, ol {{
    margin: 8pt 0;
    padding-left: 24pt;
}}
li {{
    margin: 4pt 0;
}}
hr {{
    border: none;
    border-top: 1px solid #ddd;
    margin: 16pt 0;
}}
img {{
    max-width: 100%;
}}
.toc {{
    background-color: #f9f9f9;
    border: 1px solid #ddd;
    padding: 15pt;
    margin: 20pt 0;
    border-radius: 5pt;
}}
.toc ul {{
    list-style-type: none;
    padding-left: 15pt;
}}
.toc a {{
    text-decoration: none;
}}
"""


def extract_headings(markdown_text: str) -> list:
    """Extract headings from markdown text for TOC generation."""
    headings = []
    lines = markdown_text.split('\n')
    for line in lines:
        match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
        if match:
            level = len(match.group(1))
            title = match.group(2)
            slug = re.sub(r'[^\w\s-]', '', title.lower())
            slug = re.sub(r'[-\s]+', '-', slug).strip('-')
            headings.append({'level': level, 'title': title, 'slug': slug})
    return headings


def generate_toc_html(headings: list) -> str:
    """Generate HTML table of contents."""
    if not headings:
        return ""
    
    toc_html = '<div class="toc"><h3>📑 Table des matières</h3><ul>'
    for h in headings:
        indent = (h['level'] - 1) * 20
        toc_html += f'<li style="margin-left: {indent}pt;"><a href="#{h["slug"]}">{h["title"]}</a></li>'
    toc_html += '</ul></div>'
    return toc_html


def add_heading_ids(html_content: str, headings: list) -> str:
    """Add IDs to headings for TOC linking."""
    for h in headings:
        pattern = f'<h{h["level"]}>({re.escape(h["title"])})</h{h["level"]}>'
        replacement = f'<h{h["level"]} id="{h["slug"]}">{h["title"]}</h{h["level"]}>'
        html_content = re.sub(pattern, replacement, html_content, count=1)
    return html_content


def convert_markdown_to_html(markdown_text: str) -> str:
    """Convert markdown text to HTML."""
    # Use basic extensions only to avoid rendering issues
    return markdown.markdown(
        markdown_text, 
        extensions=['tables', 'fenced_code', 'nl2br']
    )


def markdown_to_pdf(markdown_text: str, theme: dict, page_size: str, orientation: str,
                    margins: dict, show_page_numbers: bool, header_text: str, 
                    footer_text: str, generate_toc: bool, watermark: str = "") -> bytes:
    """Convert markdown text to PDF bytes using xhtml2pdf."""
    html_content = convert_markdown_to_html(markdown_text)
    
    # Generate TOC if requested
    toc_html = ""
    if generate_toc:
        headings = extract_headings(markdown_text)
        toc_html = generate_toc_html(headings)
        html_content = add_heading_ids(html_content, headings)
    
    css = generate_css(theme, page_size, orientation, margins, 
                       show_page_numbers, header_text, footer_text, watermark)
    
    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>{css}</style>
</head>
<body>
    {toc_html}
    {html_content}
</body>
</html>"""
    
    pdf_buffer = BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=pdf_buffer, encoding='utf-8')
    
    if pisa_status.err:
        raise RuntimeError(f"PDF generation failed with {pisa_status.err} errors")
    
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()


def markdown_to_html_export(markdown_text: str, theme: dict, page_size: str, 
                            orientation: str, margins: dict, generate_toc: bool) -> str:
    """Convert markdown to standalone HTML with embedded CSS."""
    html_content = convert_markdown_to_html(markdown_text)
    
    # Generate TOC if requested
    toc_html = ""
    if generate_toc:
        headings = extract_headings(markdown_text)
        toc_html = generate_toc_html(headings)
        html_content = add_heading_ids(html_content, headings)
    
    css = generate_css(theme, page_size, orientation, margins, False, "", "", "")
    
    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document exporté</title>
    <style>
        body {{
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        {css}
    </style>
</head>
<body>
    {toc_html}
    {html_content}
</body>
</html>"""


def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown to DOCX format."""
    doc = Document()
    
    lines = markdown_text.split('\n')
    in_code_block = False
    code_content = []
    in_list = False
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_content:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_content))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(10)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # Handle headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                title = match.group(2)
                # Remove markdown formatting
                title = re.sub(r'\*\*(.+?)\*\*', r'\1', title)
                title = re.sub(r'\*(.+?)\*', r'\1', title)
                heading = doc.add_heading(title, level=min(level, 9))
                continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
            continue
        
        # Handle blockquotes
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            para = doc.add_paragraph(quote_text)
            para.paragraph_format.left_indent = Inches(0.5)
            para.italic = True
            continue
        
        # Handle unordered lists
        if re.match(r'^[\s]*[-*+]\s+', line):
            list_text = re.sub(r'^[\s]*[-*+]\s+', '', line)
            # Remove markdown formatting
            list_text = re.sub(r'\*\*(.+?)\*\*', r'\1', list_text)
            list_text = re.sub(r'\*(.+?)\*', r'\1', list_text)
            doc.add_paragraph(list_text, style='List Bullet')
            continue
        
        # Handle ordered lists
        if re.match(r'^[\s]*\d+\.\s+', line):
            list_text = re.sub(r'^[\s]*\d+\.\s+', '', line)
            list_text = re.sub(r'\*\*(.+?)\*\*', r'\1', list_text)
            list_text = re.sub(r'\*(.+?)\*', r'\1', list_text)
            doc.add_paragraph(list_text, style='List Number')
            continue
        
        # Handle regular paragraphs
        if line.strip():
            para = doc.add_paragraph()
            # Process inline formatting
            text = line
            # Bold
            parts = re.split(r'\*\*(.+?)\*\*', text)
            is_bold = False
            for i, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    if i % 2 == 1:  # Bold parts
                        run.bold = True
    
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()


def fetch_markdown_from_url(url: str) -> str:
    """Fetch markdown content from a URL."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        return response.text
    except requests.RequestException as e:
        raise RuntimeError(f"Erreur lors du téléchargement: {str(e)}")


def count_stats(text: str) -> dict:
    """Count words, characters, and lines in text."""
    words = len(text.split())
    chars = len(text)
    chars_no_spaces = len(text.replace(' ', '').replace('\n', ''))
    lines = len(text.split('\n'))
    return {
        'words': words,
        'chars': chars,
        'chars_no_spaces': chars_no_spaces,
        'lines': lines
    }


def process_batch_files(files: list, theme: dict, page_size: str, orientation: str,
                        margins: dict, show_page_numbers: bool, header_text: str,
                        footer_text: str, generate_toc: bool, output_format: str) -> bytes:
    """Process multiple markdown files and return a zip archive."""
    zip_buffer = BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for uploaded_file in files:
            try:
                content = uploaded_file.read().decode('utf-8')
                base_name = uploaded_file.name.rsplit('.', 1)[0]
                
                if output_format == "PDF":
                    output_bytes = markdown_to_pdf(
                        content, theme, page_size, orientation, margins,
                        show_page_numbers, header_text, footer_text, generate_toc
                    )
                    zip_file.writestr(f"{base_name}.pdf", output_bytes)
                elif output_format == "HTML":
                    output_html = markdown_to_html_export(
                        content, theme, page_size, orientation, margins, generate_toc
                    )
                    zip_file.writestr(f"{base_name}.html", output_html.encode('utf-8'))
                elif output_format == "DOCX":
                    output_bytes = markdown_to_docx(content)
                    zip_file.writestr(f"{base_name}.docx", output_bytes)
                    
                # Reset file pointer for potential re-read
                uploaded_file.seek(0)
            except Exception as e:
                st.error(f"Erreur avec {uploaded_file.name}: {str(e)}")
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def init_session_state():
    """Initialize session state variables."""
    if 'export_history' not in st.session_state:
        st.session_state['export_history'] = []
    if 'pdf_bytes' not in st.session_state:
        st.session_state['pdf_bytes'] = None
    if 'html_bytes' not in st.session_state:
        st.session_state['html_bytes'] = None
    if 'docx_bytes' not in st.session_state:
        st.session_state['docx_bytes'] = None


def add_to_history(filename: str, format_type: str):
    """Add export to history."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state['export_history'].insert(0, {
        'filename': filename,
        'format': format_type,
        'timestamp': timestamp
    })
    # Keep only last 10 exports
    st.session_state['export_history'] = st.session_state['export_history'][:10]


def main():
    init_session_state()
    
    st.title("📄 Markdown to PDF Converter")
    st.markdown("Convertissez vos documents Markdown en PDF, HTML ou DOCX")
    
    # Sidebar with all options
    with st.sidebar:
        # Tab-like navigation in sidebar
        sidebar_section = st.radio(
            "Section",
            ["📖 Guide", "📐 Options PDF", "🎨 Thème", "📁 Fichiers", "📚 Aide Markdown"],
            label_visibility="collapsed"
        )
        
        st.divider()
        
        if sidebar_section == "📖 Guide":
            st.header("📖 Comment utiliser")
            st.markdown("""
            1. **Importer** un fichier Markdown ou coller du contenu
            2. **Configurer** les options PDF (format, marges, thème)
            3. **Prévisualiser** le rendu en temps réel
            4. **Exporter** vers PDF, HTML ou DOCX
            """)
            
            st.divider()
            st.subheader("⌨️ Raccourcis")
            st.markdown("""
            - `Ctrl+A` : Tout sélectionner
            - `Ctrl+C` : Copier
            - `Ctrl+V` : Coller
            - `Ctrl+Z` : Annuler
            - `Ctrl+Y` : Rétablir
            """)
        
        elif sidebar_section == "📐 Options PDF":
            st.header("📐 Options de page")
            
            page_size = st.selectbox(
                "Format de page",
                list(PAGE_SIZES.keys()),
                index=0
            )
            
            orientation = st.radio(
                "Orientation",
                ["Portrait", "Paysage"],
                horizontal=True
            )
            
            st.subheader("📏 Marges (cm)")
            col_m1, col_m2 = st.columns(2)
            margin_top = col_m1.number_input("Haut", 0.5, 5.0, 2.0, 0.25)
            margin_bottom = col_m2.number_input("Bas", 0.5, 5.0, 2.0, 0.25)
            margin_left = col_m1.number_input("Gauche", 0.5, 5.0, 2.0, 0.25)
            margin_right = col_m2.number_input("Droite", 0.5, 5.0, 2.0, 0.25)
            
            st.divider()
            st.subheader("📑 En-tête / Pied de page")
            show_page_numbers = st.checkbox("Numéros de page", value=True)
            header_text = st.text_input("Texte d'en-tête", placeholder="Ex: Mon Document")
            footer_text = st.text_input("Texte de pied de page", placeholder="Ex: Confidentiel")
            
            st.divider()
            st.subheader("📋 Table des matières")
            generate_toc = st.checkbox("Générer la table des matières", value=False)
            
            st.divider()
            st.subheader("💧 Filigrane")
            watermark = st.text_input("Texte du filigrane", placeholder="Ex: BROUILLON")
        
        elif sidebar_section == "🎨 Thème":
            st.header("🎨 Style du document")
            
            selected_theme = st.selectbox(
                "Thème",
                list(THEMES.keys()),
                index=0
            )
            
            st.divider()
            st.subheader("👁️ Aperçu")
            show_preview = st.checkbox("Afficher la prévisualisation", value=True)
            preview_height = st.slider("Hauteur de prévisualisation", 300, 800, 500, 50)
            
            # Theme preview
            theme = THEMES[selected_theme]
            st.markdown(f"""
            <div style="padding: 10px; border: 1px solid #ddd; border-radius: 5px; font-size: 12px;">
                <p style="color: {theme['heading_color']}; font-weight: bold;">Titre exemple</p>
                <p style="color: {theme['text_color']};">Texte normal</p>
                <p><a style="color: {theme['link_color']};">Lien exemple</a></p>
                <code style="background: {theme['code_bg']}; padding: 2px 4px;">code</code>
            </div>
            """, unsafe_allow_html=True)
        
        elif sidebar_section == "📁 Fichiers":
            st.header("📁 Gestion des fichiers")
            
            st.subheader("📥 Import depuis URL")
            url_input = st.text_input("URL du fichier .md", placeholder="https://...")
            if st.button("📥 Charger depuis URL", use_container_width=True):
                if url_input:
                    try:
                        content_from_url = fetch_markdown_from_url(url_input)
                        st.session_state['url_content'] = content_from_url
                        st.success("Contenu chargé avec succès!")
                    except Exception as e:
                        st.error(str(e))
            
            st.divider()
            st.subheader("📦 Conversion par lot")
            batch_files = st.file_uploader(
                "Sélectionner plusieurs fichiers",
                type=['md', 'markdown', 'txt'],
                accept_multiple_files=True,
                key="batch_uploader"
            )
            
            if batch_files:
                st.info(f"{len(batch_files)} fichier(s) sélectionné(s)")
                batch_format = st.selectbox("Format de sortie", ["PDF", "HTML", "DOCX"])
                
                if st.button("🔄 Convertir tous", use_container_width=True):
                    # Get current settings from session or defaults
                    with st.spinner("Conversion en cours..."):
                        try:
                            zip_bytes = process_batch_files(
                                batch_files,
                                THEMES.get(st.session_state.get('selected_theme', 'Professionnel'), THEMES['Professionnel']),
                                st.session_state.get('page_size', 'A4'),
                                st.session_state.get('orientation', 'Portrait'),
                                st.session_state.get('margins', {'top': 2, 'bottom': 2, 'left': 2, 'right': 2}),
                                st.session_state.get('show_page_numbers', True),
                                st.session_state.get('header_text', ''),
                                st.session_state.get('footer_text', ''),
                                st.session_state.get('generate_toc', False),
                                batch_format
                            )
                            st.session_state['batch_zip'] = zip_bytes
                            st.success("Conversion terminée!")
                        except Exception as e:
                            st.error(f"Erreur: {str(e)}")
                
                if 'batch_zip' in st.session_state:
                    st.download_button(
                        "⬇️ Télécharger le ZIP",
                        data=st.session_state['batch_zip'],
                        file_name="converted_files.zip",
                        mime="application/zip",
                        use_container_width=True
                    )
            
            st.divider()
            st.subheader("📜 Historique des exports")
            if st.session_state['export_history']:
                for item in st.session_state['export_history'][:5]:
                    st.text(f"📄 {item['filename']}.{item['format'].lower()}")
                    st.caption(f"   {item['timestamp']}")
            else:
                st.caption("Aucun export récent")
        
        elif sidebar_section == "📚 Aide Markdown":
            st.header("📚 Aide Markdown")
            st.markdown(MARKDOWN_CHEATSHEET, unsafe_allow_html=True)
    
    # Store sidebar settings in session state for batch processing
    if sidebar_section == "📐 Options PDF":
        st.session_state['page_size'] = page_size
        st.session_state['orientation'] = orientation
        st.session_state['margins'] = {
            'top': margin_top, 'bottom': margin_bottom,
            'left': margin_left, 'right': margin_right
        }
        st.session_state['show_page_numbers'] = show_page_numbers
        st.session_state['header_text'] = header_text
        st.session_state['footer_text'] = footer_text
        st.session_state['generate_toc'] = generate_toc
        st.session_state['watermark'] = watermark
    elif sidebar_section == "🎨 Thème":
        st.session_state['selected_theme'] = selected_theme
        st.session_state['show_preview'] = show_preview
        st.session_state['preview_height'] = preview_height
    
    # Get current settings with defaults
    current_page_size = st.session_state.get('page_size', 'A4')
    current_orientation = st.session_state.get('orientation', 'Portrait')
    current_margins = st.session_state.get('margins', {'top': 2, 'bottom': 2, 'left': 2, 'right': 2})
    current_show_page_numbers = st.session_state.get('show_page_numbers', True)
    current_header_text = st.session_state.get('header_text', '')
    current_footer_text = st.session_state.get('footer_text', '')
    current_generate_toc = st.session_state.get('generate_toc', False)
    current_watermark = st.session_state.get('watermark', '')
    current_theme_name = st.session_state.get('selected_theme', 'Professionnel')
    current_theme = THEMES[current_theme_name]
    current_show_preview = st.session_state.get('show_preview', True)
    current_preview_height = st.session_state.get('preview_height', 500)
    
    # Main content area with two columns
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("✏️ Éditeur Markdown")
        
        # File upload option
        uploaded_file = st.file_uploader(
            "📂 Importer un fichier Markdown",
            type=['md', 'markdown', 'txt'],
            key="single_uploader"
        )
        
        # Default content
        default_content = """# Bienvenue dans le Convertisseur Markdown

## Fonctionnalités

- **Facile à utiliser** : Collez ou importez votre markdown
- **Prévisualisation en direct** : Voyez les changements en temps réel
- **Export multi-format** : PDF, HTML ou DOCX

## Exemple de code

```python
def hello_world():
    print("Bonjour, le monde!")
```

## Tableau d'exemple

| Fonctionnalité | Statut |
|----------------|--------|
| Import fichier | ✅ |
| Édition live | ✅ |
| Export PDF | ✅ |
| Export HTML | ✅ |
| Export DOCX | ✅ |

## Liste ordonnée

1. Premier élément
2. Deuxième élément
3. Troisième élément

> Ceci est un exemple de citation

---

**Texte en gras** et *texte en italique*
"""
        
        # Priority: URL content > uploaded file > default
        if 'url_content' in st.session_state:
            content = st.session_state['url_content']
            del st.session_state['url_content']  # Clear after use
        elif uploaded_file is not None:
            try:
                content = uploaded_file.read().decode('utf-8')
            except UnicodeDecodeError:
                st.error("Erreur: Impossible de décoder le fichier. Veuillez utiliser un fichier texte valide.")
                content = default_content
        else:
            content = default_content
        
        # Text area for editing
        markdown_text = st.text_area(
            "Éditez votre Markdown ici:",
            value=content,
            height=current_preview_height,
            key="markdown_editor"
        )
        
        # Statistics
        stats = count_stats(markdown_text)
        st.caption(
            f"📊 {stats['words']} mots | {stats['chars']} caractères | "
            f"{stats['chars_no_spaces']} sans espaces | {stats['lines']} lignes"
        )
    
    with col2:
        if current_show_preview:
            st.subheader("👁️ Prévisualisation")
            
            try:
                html_content = convert_markdown_to_html(markdown_text)
                
                # Add TOC if enabled
                if current_generate_toc:
                    headings = extract_headings(markdown_text)
                    toc_html = generate_toc_html(headings)
                    html_content = toc_html + add_heading_ids(html_content, headings)
                
                # Generate preview CSS
                preview_css = generate_css(
                    current_theme, current_page_size, 'Portrait',
                    {'top': 0, 'bottom': 0, 'left': 0, 'right': 0},
                    False, '', '', ''
                )
                
                # Use components.html for proper CSS rendering
                full_preview_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <style>
                        body {{
                            margin: 0;
                            padding: 20px;
                            background: white;
                        }}
                        {preview_css}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                components.html(full_preview_html, height=current_preview_height, scrolling=True)
            except Exception as e:
                st.error(f"Erreur de rendu: {str(e)}")
        else:
            st.info("Activez la prévisualisation dans les options Thème de la sidebar")
    
    # Export section
    st.divider()
    
    # Filename input
    col_name1, col_name2 = st.columns([3, 1])
    with col_name1:
        export_filename = st.text_input(
            "📝 Nom du fichier (sans extension)",
            value="document",
            key="export_filename"
        )
    
    # Export buttons
    st.markdown("### 📤 Exporter le document")
    col_exp1, col_exp2, col_exp3, col_exp4 = st.columns(4)
    
    with col_exp1:
        if st.button("📕 Générer PDF", use_container_width=True):
            try:
                with st.spinner("Génération du PDF..."):
                    pdf_bytes = markdown_to_pdf(
                        markdown_text, current_theme, current_page_size,
                        current_orientation, current_margins, current_show_page_numbers,
                        current_header_text, current_footer_text, current_generate_toc,
                        current_watermark
                    )
                    st.session_state['pdf_bytes'] = pdf_bytes
                    add_to_history(export_filename, 'PDF')
                    st.success("PDF généré!")
            except Exception as e:
                st.error(f"Erreur: {str(e)}")
    
    with col_exp2:
        if st.button("🌐 Générer HTML", use_container_width=True):
            try:
                with st.spinner("Génération du HTML..."):
                    html_export = markdown_to_html_export(
                        markdown_text, current_theme, current_page_size,
                        current_orientation, current_margins, current_generate_toc
                    )
                    st.session_state['html_bytes'] = html_export.encode('utf-8')
                    add_to_history(export_filename, 'HTML')
                    st.success("HTML généré!")
            except Exception as e:
                st.error(f"Erreur: {str(e)}")
    
    with col_exp3:
        if st.button("📘 Générer DOCX", use_container_width=True):
            try:
                with st.spinner("Génération du DOCX..."):
                    docx_bytes = markdown_to_docx(markdown_text)
                    st.session_state['docx_bytes'] = docx_bytes
                    add_to_history(export_filename, 'DOCX')
                    st.success("DOCX généré!")
            except Exception as e:
                st.error(f"Erreur: {str(e)}")
    
    with col_exp4:
        if st.button("🔄 Réinitialiser", use_container_width=True):
            for key in ['pdf_bytes', 'html_bytes', 'docx_bytes']:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()
    
    # Download buttons
    st.markdown("### ⬇️ Télécharger")
    col_dl1, col_dl2, col_dl3, col_dl4 = st.columns(4)
    
    with col_dl1:
        if st.session_state.get('pdf_bytes'):
            st.download_button(
                label="⬇️ PDF",
                data=st.session_state['pdf_bytes'],
                file_name=f"{export_filename}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        else:
            st.button("⬇️ PDF", disabled=True, use_container_width=True)
    
    with col_dl2:
        if st.session_state.get('html_bytes'):
            st.download_button(
                label="⬇️ HTML",
                data=st.session_state['html_bytes'],
                file_name=f"{export_filename}.html",
                mime="text/html",
                use_container_width=True
            )
        else:
            st.button("⬇️ HTML", disabled=True, use_container_width=True)
    
    with col_dl3:
        if st.session_state.get('docx_bytes'):
            st.download_button(
                label="⬇️ DOCX",
                data=st.session_state['docx_bytes'],
                file_name=f"{export_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.button("⬇️ DOCX", disabled=True, use_container_width=True)
    
    with col_dl4:
        # Copy HTML to clipboard (simulated with text area)
        if st.session_state.get('html_bytes'):
            if st.button("📋 Copier HTML", use_container_width=True):
                st.code(st.session_state['html_bytes'].decode('utf-8')[:500] + "...", language="html")
                st.info("Sélectionnez et copiez le code HTML ci-dessus")


if __name__ == "__main__":
    main()
